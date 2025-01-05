import pytesseract
from pdf2image import convert_from_path
import cv2
import numpy as np
from PIL import Image
import pandas as pd
from pathlib import Path
import logging
from PyPDF2 import PdfReader
import docx
import os

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.tiff': self._process_image,
            '.xlsx': self._process_excel,
            '.csv': self._process_csv,
            '.docx': self._process_docx
        }

    def process(self, file_path: Path):
        """Main processing method that handles different file types"""
        try:
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")

            processor = self.supported_formats[file_extension]
            return processor(file_path)

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}", exc_info=True)
            raise

    def _process_pdf(self, file_path: Path):
        """Process PDF files using both text extraction and OCR"""
        try:
            # Extract text using PyPDF2
            reader = PdfReader(str(file_path))
            text_content = []
            images = []

            for page_num, page in enumerate(reader.pages):
                # Extract text
                text = page.extract_text()
                text_content.append({"page": page_num + 1, "text": text})

                # Convert page to image for OCR and table detection
                pdf_images = convert_from_path(str(file_path), first_page=page_num+1, last_page=page_num+1)
                for img in pdf_images:
                    img_array = np.array(img)
                    # Perform OCR if text extraction yielded little content
                    if len(text.strip()) < 100:
                        ocr_text = pytesseract.image_to_string(img_array)
                        text_content.append({"page": page_num + 1, "ocr_text": ocr_text})
                    
                    # Store image for table detection
                    images.append({"page": page_num + 1, "image": img_array})

            return {
                "type": "pdf",
                "text_content": text_content,
                "images": images
            }

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}", exc_info=True)
            raise

    def _process_image(self, file_path: Path):
        """Process image files with advanced OCR and table detection"""
        try:
            # Read image
            image = cv2.imread(str(file_path))
            if image is None:
                raise ValueError(f"Could not read image file: {file_path}")

            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

            # Apply preprocessing for better OCR
            denoised = cv2.fastNlMeansDenoising(gray)
            _, binary = cv2.threshold(denoised, 128, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

            # Perform OCR with additional configuration
            ocr_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(binary, config=ocr_config)

            # Get bounding boxes for text regions
            boxes = pytesseract.image_to_data(binary, output_type=pytesseract.Output.DICT)

            return {
                "type": "image",
                "text": text,
                "boxes": boxes,
                "image": image
            }

        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}", exc_info=True)
            raise

    def _process_excel(self, file_path: Path):
        """Process Excel files"""
        try:
            df = pd.read_excel(str(file_path))
            return {
                "type": "excel",
                "data": df.to_dict(orient="records"),
                "columns": df.columns.tolist(),
                "shape": df.shape
            }
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {str(e)}", exc_info=True)
            raise

    def _process_csv(self, file_path: Path):
        """Process CSV files"""
        try:
            df = pd.read_csv(str(file_path))
            return {
                "type": "csv",
                "data": df.to_dict(orient="records"),
                "columns": df.columns.tolist(),
                "shape": df.shape
            }
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {str(e)}", exc_info=True)
            raise

    def _process_docx(self, file_path: Path):
        """Process Word documents"""
        try:
            doc = docx.Document(str(file_path))
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append({
                        "type": "paragraph",
                        "text": paragraph.text
                    })

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content.append({
                    "type": "table",
                    "data": table_data
                })

            return {
                "type": "docx",
                "content": content
            }
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}", exc_info=True)
            raise
